
#pip install python-docx

from docx import Document
from docx.shared import Inches
import csv

# Find and replace the Keywords in the File
old_name = 'AADHAR HOSPITAL'
old_address = 'Near, South Bypass Crossing NH10, Hisar, Haryana 125001'
oldRates = 'MYRATES'


def makeword(name,rate,address):
    # Open the existing Word document
    doc = Document('Sample File.docx')


    # Create a new document as a clone of the original document
    new_doc = Document()
    for element in doc.element.body:
        new_doc.element.body.append(element)
 
    new_name = name
    new_address = address
    newrates = rate
    target_text = "SignaturesofYourOrganization"
    targetrates = oldRates
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(old_name, new_name).replace(old_address, new_address)
            run.text = run.text.replace(oldRates, newrates)
            if target_text in run.text:
                run.text = run.text.replace(target_text, '')
                run.add_picture('Sign.png', width=Inches(2))
                print('Quotation Submit')

    # Save the modified document with a new name
    new_doc.save(f'{name}.docx')

try:
    # Open the CSV file
    with open('List.csv', 'r') as file:
        # Create a CSV reader object
        reader = csv.DictReader(file)
        # Iterate over each row in the CSV file
        for row in reader:
            # Store the values of Names, Rates, and Address in variables
            name = row['Names']
            rate = row['Rates']
            address = row['Address']
            # Print the values
            print("Name:", name)
            makeword(name,rate,address)
            print()
    input('Press Any key to Exit')

except Exception as e:
    print(e)
    input('Press Any key to exit')
